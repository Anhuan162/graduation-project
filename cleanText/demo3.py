import re
from docx import Document

# ================================
#   DICTIONARY SỬA LỖI OCR
# ================================
replacement_dict = {
    "thong tin": "thông tin",
    "hoc vien": "học viện",
    "cong nghe": "công nghệ",
    "buư chính": "bưu chính",
    "vien thong": "viễn thông",
    "qd hv": "qđ-hv",
    "cong hoa": "cộng hòa",
    "xa hoi": "xã hội",
    "doc lap": "độc lập",
    "tu do": "tự do",
    "hanh phuc": "hạnh phúc",
}

# ================================
#  HÀM GHÉP TỪ BỊ TÁCH CHỮ: "h ọ c v i ệ n" → "học viện"
# ================================
def fix_broken_words(text):
    # trường hợp OCR tách từng chữ → ghép lại
    text = re.sub(r"(\b(?:[a-zA-ZÀ-Ỵà-ỹ]\s){1,6}[a-zA-ZÀ-Ỵà-ỹ]\b)",
                  lambda m: m.group(1).replace(" ", ""),
                  text)
    return text

# ================================
#  HÀM LÀM SẠCH OCR
# ================================
def clean_text(text):

    # 1. Đưa toàn bộ về chữ thường
    text = text.lower()

    # 2. Xóa ký tự rác không thuộc tiếng Việt
    text = re.sub(r"[^0-9a-zÀ-Ỵà-ỹ.,?!:;/()\-\n\s]", " ", text)

    # 3. Xóa dãy ký tự lặp (---, +++, ***)
    text = re.sub(r"[-+*_/]{2,}", " ", text)

    # 4. Ghép dòng bị gãy (xuống dòng giữa câu)
    text = re.sub(r"\n(?=[a-zà-ỹ0-9])", " ", text)

    # 5. Sửa lỗi tách chữ OCR
    text = fix_broken_words(text)

    # 6. Chuẩn hóa khoảng trắng
    text = re.sub(r"\s+", " ", text)

    # 7. Đưa xuống dòng trước "chương", "điều"
    text = re.sub(r"\s*(chương\s+[ivxlcdm]+)", r"\n\n\1\n", text, flags=re.I)
    text = re.sub(r"\s*(điều\s+\d+)", r"\n\n\1\n", text, flags=re.I)

    # 8. Chuẩn hóa dấu câu
    text = re.sub(r"\s*\.\s*\.", ".", text)
    text = re.sub(r"\s*,\s*,", ",", text)
    text = re.sub(r"\s*\.\s*,", ".", text)

    # 9. Sửa OCR dựa trên dictionary
    for wrong, correct in replacement_dict.items():
        text = re.sub(rf"\b{wrong}\b", correct, text)

    return text.strip()


# ================================
#  HÀM LÀM SẠCH DOCX
# ================================
def clean_docx(input_path, output_path):
    doc = Document(input_path)

    # gộp toàn bộ text
    raw = "\n".join(p.text for p in doc.paragraphs)

    cleaned = clean_text(raw)

    # lưu file kết quả
    new_doc = Document()
    for line in cleaned.split("\n"):
        new_doc.add_paragraph(line)

    new_doc.save(output_path)
    print("Đã làm sạch thành công →", output_path)


# ================================
#   CHẠY TOOL
# ================================
input_file = "test2.docx"
output_file = "output_clean_2.docx"

clean_docx(input_file, output_file)
