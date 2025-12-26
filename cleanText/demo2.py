import re
from docx import Document
# Danh sách các cụm từ OCR hay bị tách trong tài liệu hành chính VN
OCR_FIXES = {
    r"B Ộ": "BỘ",
    r"T H Ô N G": "THÔNG",
    r"V I Ễ N": "VIỄN",
    r"C Ô N G": "CÔNG",
    r"N G H Ệ": "NGHỆ",
    r"B Ư U": "BƯU",
    r"C H Í N H": "CHÍNH",
    r"T R U Y Ề N": "TRUYỀN",

    r"H Ọ C\s+V I Ệ N": "HỌC VIỆN",
    r"C Ô N G\s+N G H Ệ": "CÔNG NGHỆ",
    r"B Ư U\s+C H Í N H": "BƯU CHÍNH",
}

# --------- HÀM LÀM SẠCH NÂNG CẤP -----------
def clean_text(text):
    if not text:
        return ""

    # 1. Xóa ký tự lạ
    text = re.sub(r"[^0-9A-Za-zÀ-Ỵà-ỵ.,?!:;/()\-\n\s]", " ", text)

    # 2. Chuẩn hóa khoảng trắng
    text = re.sub(r"\s+", " ", text)

    # 3. Ghép dòng bị tách giữa câu (nếu dòng dưới bắt đầu bằng chữ thường)
    text = re.sub(r"\n(?=[a-zà-ỹ])", " ", text)

    # 4. Ghép dòng nếu dòng trước không kết thúc bằng dấu câu
    text = re.sub(r"(?<![.!?;:])\n(?=[A-ZÀ-Ỵ])", " ", text)

    # 5. Ghép từ bị tách bằng khoảng trắng giữa mỗi ký tự
    text = re.sub(r"(?<!\w)([A-Za-zÀ-Ỵà-ỹ])\s+(?=[A-Za-zÀ-Ỵà-ỹ](\s|$))", r"\1", text)

    # 6. Áp dụng sửa lỗi OCR bằng từ điển
    for pattern, replacement in OCR_FIXES.items():
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    # 7. Xóa dấu chấm câu sai
    text = re.sub(r"\s*\.\s*\.", ".", text)
    text = re.sub(r"\s*,\s*,", ",", text)

    return text.strip()


# --------- TOOL LÀM SẠCH DOCX + GIỮ BẢNG -----------
def clean_docx(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    # ---- Xử lý Paragraph ----
    for para in doc.paragraphs:
        cleaned_para = clean_text(para.text)
        if cleaned_para.strip():
            new_doc.add_paragraph(cleaned_para)

    # ---- Xử lý TABLE ----
    for table in doc.tables:
        # Tạo bảng mới trong file output
        new_table = new_doc.add_table(rows=0, cols=len(table.columns))
        new_table.style = "Table Grid"

        for row in table.rows:
            new_row = new_table.add_row().cells
            for i, cell in enumerate(row.cells):
                cleaned_cell = clean_text(cell.text)
                new_row[i].text = cleaned_cell

    # Lưu file
    new_doc.save(output_path)
    print("Đã làm sạch và giữ bảng. Xuất file:", output_path)


# --------- CHẠY TOOL -----------
input_file = "1573qd.docx"
output_file = "output_clean42.docx"

clean_docx(input_file, output_file)
